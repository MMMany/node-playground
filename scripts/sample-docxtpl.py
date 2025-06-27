from argparse import ArgumentParser
from docxtpl import DocxTemplate
import os
from pathlib import Path
from docx import Document
import re


def get_bookmarked_tables(doc):
    xml = doc._element.xml
    pattern = r'<w:bookmarkStart w:id="[0-9]+" w:name="([^"]+)"'
    bookmarks = re.findall(pattern, xml)
    print("bookmarks :", bookmarks)
    result = {}
    for table in doc.tables:
        table_xml = table._element.xml
        for name in bookmarks:
            if table_xml.find(name) > 0:
                result[name] = table
                break
    return result


def merge_table_cells_by_row(table, target_table, data):
    target_cols = data[target_table]["merged_cols"]
    current_category = None
    start_row = None

    for col_idx in target_cols:
        for i in range(1, len(table.rows)):
            cell_value = table.cell(i, col_idx).text.strip()

            if cell_value != current_category:
                if start_row is not None and i > start_row + 1:
                    for row_idx in range(start_row + 1, i):
                        cell = table.cell(row_idx, col_idx)
                        for paragraph in cell.paragraphs:
                            paragraph.clear()

                    table.cell(start_row, col_idx).merge(table.cell(i - 1, col_idx))

                current_category = cell_value
                start_row = i

        if start_row is not None and start_row < len(table.rows) - 1:
            for row_idx in range(start_row + 1, len(table.rows)):
                cell = table.cell(row_idx, col_idx)
                for paragraph in cell.paragraphs:
                    paragraph.clear()

            table.cell(start_row, col_idx).merge(
                table.cell(len(table.rows) - 1, col_idx)
            )


def merge_table_cells(table, target_table, data):
    # if not table or target_table not in data.keys():
    #     return

    print(f"Found table: {target_table}")
    target_cols = data[target_table]["merged_cols"]

    # 1. 셀 값 수집 및 병합 영역 식별
    areas_to_merge = []  # [(start_row, start_col, end_row, end_col), ...]

    # 헤더를 제외한 행들에 대해
    row_count = len(table.rows)
    col_count = len(table.columns)

    # 1.1 값이 같은 영역 식별 (target_cols 내의 인접한 영역만)
    for col_idx in target_cols:
        for row_idx in range(1, row_count):  # 헤더 제외
            if not is_valid_cell(row_idx, col_idx, table):
                continue

            cell_value = table.cell(row_idx, col_idx).text.strip()
            if not cell_value:  # 빈 셀은 건너뛰기
                continue

            # 이미 병합 영역에 포함된 셀인지 확인
            already_merged = False
            for area in areas_to_merge:
                sr, sc, er, ec = area
                if sr <= row_idx <= er and sc <= col_idx <= ec:
                    already_merged = True
                    break
            if already_merged:
                continue

            # 최대 병합 영역 찾기
            max_row = row_idx
            max_col = col_idx

            # 행 방향으로 확장
            for r in range(row_idx + 1, row_count):
                if not is_valid_cell(r, col_idx, table):
                    break
                if table.cell(r, col_idx).text.strip() != cell_value:
                    break
                max_row = r

            # 열 방향으로 확장 (target_cols 안의 값들만 고려)
            next_col_idx = get_next_target_col(col_idx, target_cols)
            while next_col_idx < col_count and next_col_idx != -1:
                can_extend = True

                # 확장된 열에 대해 모든 행이 같은 값을 가지는지 확인
                for r in range(row_idx, max_row + 1):
                    if (
                        not is_valid_cell(r, next_col_idx, table)
                        or table.cell(r, next_col_idx).text.strip() != cell_value
                    ):
                        can_extend = False
                        break

                if not can_extend:
                    break

                max_col = next_col_idx
                next_col_idx = get_next_target_col(next_col_idx, target_cols)

            # 영역이 2×2 이상이면 병합 대상에 추가
            if max_row > row_idx or max_col > col_idx:
                areas_to_merge.append((row_idx, col_idx, max_row, max_col))

    # 2. 병합 실행
    for area in areas_to_merge:
        start_row, start_col, end_row, end_col = area
        try:
            # 병합할 영역의 시작점을 제외한 모든 셀 내용 지우기
            for r in range(start_row, end_row + 1):
                for c in range(start_col, end_col + 1):
                    if r != start_row or c != start_col:
                        cell = table.cell(r, c)
                        for paragraph in cell.paragraphs:
                            paragraph.clear()

            # 병합 실행 (직사각형 영역 전체를 한번에 병합)
            table.cell(start_row, start_col).merge(table.cell(end_row, end_col))
            print(
                f"영역 병합 완료: ({start_row}, {start_col}) - ({end_row}, {end_col})"
            )
        except Exception as e:
            print(
                f"영역 병합 오류: ({start_row}, {start_col}) - ({end_row}, {end_col}): {e}"
            )


# 유효한 셀인지 확인하는 헬퍼 함수
def is_valid_cell(row, col, table):
    try:
        table.cell(row, col)
        return True
    except Exception:
        return False


# target_cols 내에서 다음 열 인덱스를 찾는 헬퍼 함수
def get_next_target_col(current_col, target_cols):
    sorted_cols = sorted(target_cols)
    for col in sorted_cols:
        if col > current_col:
            return col
    return -1


def default_run():
    template_path = Path(os.path.join(os.getcwd(), "templates/sample-docxtpl-tpl.docx"))
    template_output = Path(
        os.path.join(os.getcwd(), "output/sample-docxtpl-output.docx")
    )
    output_path = Path(os.path.join(os.getcwd(), "output/sample-pydocx-output.docx"))
    data_path = Path(os.path.join(os.getcwd(), "docxtpl-meta.json"))

    template = DocxTemplate(template_path)

    data = None
    with open(data_path, "r", encoding="utf8") as f:
        data = eval(f.read())

    assert data is not None, "failed read meta data"

    template.render(data)
    template.save(template_output)

    docx = Document(template_output)

    tables = get_bookmarked_tables(docx)
    for name, table in tables.items():
        merge_table_cells(table, name, data)

    docx.save(output_path)


def merge_cells_run(target, meta, output):
    meta_path = Path(os.path.abspath(meta))
    data = None
    with open(meta_path, "r", encoding="utf8") as f:
        data = eval(f.read())
    assert data is not None, "failed read meta data"

    target_path = Path(os.path.abspath(target))
    output_path = Path(os.path.abspath(output))

    docx = Document(target_path)
    tables = get_bookmarked_tables(docx)
    for name, table in tables.items():
        merge_table_cells(table, name, data)

    docx.save(output_path)


def main():
    parser = ArgumentParser()
    parser.add_argument("-t", "--target", type=str)
    parser.add_argument("-m", "--meta", type=str)
    parser.add_argument("-o", "--output", type=str)
    args = parser.parse_args()

    required_args = ["target", "meta", "output"]

    if [arg for arg in required_args if vars(args)[arg] is None]:
        default_run()
    else:
        merge_cells_run(args.target, args.meta, args.output)


if __name__ == "__main__":
    main()
